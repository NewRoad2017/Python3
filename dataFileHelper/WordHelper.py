from docx import  Document
import glob
import os
import sys
sys.path.append("D:/Applications/Python3/GeneralFunction")
from GeneralFunction import ReHelper,TableHelper

class Word:
    def __init__(self):
        pass
    def fileRead(self,dataPath):
        self.word = Document(dataPath)
        self.listParagraphs = self.word.paragraphs
        self.paragraphCount = len(self.listParagraphs)
        self.listTables = self.word.tables
        self.tableCount = len(self.listTables)

    def GetuniTable(self, index,dataHasHead):
        name = "table" + str(index)
        tableDOCX = self.listTables[index]
        rowsCount = len(tableDOCX.rows)
        listTableRows = []
        if dataHasHead == 0:
            for index in range(0, rowsCount):
                rowCur = tableDOCX.rows[index]
                listRowparts = self.__getlistRowText(rowCur)
                listTableRows.append(listRowparts)
            table = TableHelper.Table(name, listTableRows, [])
        else:
            for index in range(0, rowsCount):
                rowCur = tableDOCX.rows[index]
                listTableHeaditems = []
                if index == 0:
                    listTableHeaditems = self.__getlistRowText(rowCur)
                else:
                    listRowparts = self.__getlistRowText(rowCur)
                    listTableRows.append(listRowparts)
            table = TableHelper.Table(name,listTableRows,1,listTableHeaditems)
        return table

    def __getlistRowText(self, dataRowCur):
        list = []
        listRowCells = dataRowCur.cells
        for cell in listRowCells:
            cellText = self.__getAllParaInaCell(cell)
            list.append(cellText)
        return list
    def __getAllParaInaCell(self, dataCell):
        listParaTextParts = []
        str = ''
        listParagraphs = dataCell.paragraphs
        for paragraph in listParagraphs:
            listParaTextParts.append(paragraph.text)
            str = str + paragraph.text
        return str



# class WordRead:
#     def __init__(self, dataPath):
#         self.filePath = dataPath
#         self.wordDOCX = Document(self.filePath)
#         self.listParagraphs = self.wordDOCX.paragraphs
#         self.paragraphCount = len(self.listParagraphs)
#         self.listTables=self.wordDOCX.tables
#         self.tableCount = len(self.listTables)
#
#     def GetuniTable(self, index,dataHasHead):
#         name = "table" + str(index)
#         tableDOCX = self.listTables[index]
#         colsCount = len(tableDOCX.rows[0].cells)
#         rowsCount = len(tableDOCX.rows)
#         listTableRows = []
#         if dataHasHead == 0:
#             for index in range(0, rowsCount):
#                 rowCur = tableDOCX.rows[index]
#                 listRowparts = self.__getlistRowText(rowCur)
#                 listTableRows.append(listRowparts)
#             table = TableHelper.Table(name, listTableRows, [])
#         else:
#             for index in range(0, rowsCount):
#                 rowCur = tableDOCX.rows[index]
#                 listTableHeaditems = []
#                 if index == 0:
#                     listTableHeaditems = self.__getlistRowText(rowCur)
#                 else:
#                     listRowparts = self.__getlistRowText(rowCur)
#                     listTableRows.append(listRowparts)
#             table = TableHelper.Table(name,listTableRows,listTableHeaditems)
#         return table
#
#     def __getlistRowText(self, dataRowCur):
#         list = []
#         listRowCells = dataRowCur.cells
#         for cell in listRowCells:
#             cellText = self.__getAllParaInaCell(cell)
#             list.append(cellText)
#         return list
#     def __getAllParaInaCell(self, dataCell):
#         listParaTextParts = []
#         str = ''
#         listParagraphs = dataCell.paragraphs
#         for paragraph in listParagraphs:
#             listParaTextParts.append(paragraph.text)
#             str = str + paragraph.text
#         return str

if __name__ == "__main__":
    filePath = r"C:\Users\chenlu\Desktop\word Version\test\south_33_register\257.docx"
    filePath = ReHelper.replacePath(filePath)
    word = Word()
    word.fileRead(filePath)
    count = word.paragraphCount
    listParas = word.listParagraphs
    table = word.GetuniTable(0,0)
    for para in listParas:
        str = para.text
        print(str)
    print(2)
    # word.GetuniTable(0)




    # extract the first columns into a list
    # listHeads=[]
    # rowCur = table.rows[1]
    # listCells = rowCur.cells
    # countListCells = len(listCells)
    # for cell in listCells:
    #     listParagraph = cell.paragraphs
    #     text = ""
    #     for paragraph in listParagraph:
    #         text = text + paragraph.text
    #     listHeads.append(text)
    # listRowsText = []
    # listHeadsText = []
    # dictInfo = {}
    # for index in range(0,countRows):
    #     rowCur = table.rows[index]
    #     listCells = rowCur.cells
    #     countListCells = len(listCells)
    #     if index in [0,1,6,7,8,9]:
    #         cellCur = listCells[0]
    #         strHead = __getAllParaInaCell(cellCur)
    #         listHeads.append(strHead)
    #     elif index in [4,5]:
    #         for indexcell in range(0,6):
    #             cellCur = listCells[indexcell]
    #             strHead = __getAllParaInaCell(cellCur)
    #             listHeads.append(strHead)

        # for indexCell in range(0,countListCells):
        #     if index == 0:
        #         strHead = ''
        #         listHeadParagraphs = listCells[indexCell].paragraphs
        #         for paragraph in listHeadParagraphs:
        #             strHead = strHead+paragraph

                # if indexCell % 2 == 0:
                #     listHeadParagraphs = listCells[indexCell].paragraphs
                #     title = ""
                #     for paragraph in listHeadParagraphs:
                #         strHead = title+paragraph.text
                #         listHeads.append(strHead)
    # print (listHeads)
    #print (table.rows[0].cells[0].paragraphs[0].text)
    print (2)